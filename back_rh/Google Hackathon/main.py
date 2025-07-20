from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import fitz  # PyMuPDF
import os
import tempfile
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from datetime import datetime

# Charger les variables d'environnement
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-2.0-flash-001")

app = FastAPI(title="API Gemini PDF Analyzer")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_pdf(pdf_path):
    """
    Extrait le texte d'un fichier PDF à l'aide de PyMuPDF.
    """
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

# ... [tes fonctions extract_text_from_pdf, generate_questions_with_gemini, generate_criteria_from_pdf, generate_criteria_from_job_description, generate_docx ici] ...

def generate_criteria_from_pdf(pdf_path):
    """
    Extrait les critères d'un PDF en utilisant l'extraction de texte.
    Cette fonction est un exemple et doit être adaptée selon la logique métier.
    """
    text = extract_text_from_pdf(pdf_path)
    # Exemple simple : chaque ligne non vide est considérée comme un critère
    criteria = [line.strip() for line in text.split('\n') if line.strip()]
    return criteria

def generate_questions_with_gemini(pdf_path, nb_questions):
    # Extraction du texte du PDF
    text = extract_text_from_pdf(pdf_path)
    # Génération des questions avec Gemini
    prompt = f"Génère {nb_questions} questions à choix multiples techniques en te basant sur l'idée generale de :\n{text}"
    response = model.generate_content(prompt)
    # Supposons que la réponse est une liste de questions séparées par des sauts de ligne
    questions = response.text.strip().split('\n')
    return questions

def generate_docx(candidate_name, job_title, analyst_name, criteria_list, output_path):
    doc = Document()
    doc.add_heading('Fiche d\'évaluation', 0)
    doc.add_paragraph(f"Nom du candidat : {candidate_name}")
    doc.add_paragraph(f"Poste : {job_title}")
    doc.add_paragraph(f"Analyste : {analyst_name}")
    doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_heading('Critères d\'évaluation', level=1)
    for critere in criteria_list:
        doc.add_paragraph(critere, style='List Bullet')
    doc.save(output_path)

def generate_criteria_from_job_description(text):
    """
    Extrait les critères à partir d'une description de poste (texte brut).
    Cette fonction est un exemple simple à adapter selon la logique métier.
    """
    # Exemple simple : chaque ligne non vide est considérée comme un critère
    criteria = [line.strip() for line in text.split('\n') if line.strip()]
    return criteria

@app.post("/generate-questions")
async def generate_questions_endpoint(
    file: UploadFile = File(...),
    nb_questions: int = Form(5)
):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name

        questions = generate_questions_with_gemini(tmp_path, nb_questions)
        os.remove(tmp_path)

        return JSONResponse(content={"questions": questions})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/generate-criteria")
async def generate_criteria_endpoint(file: UploadFile = File(...)):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name

        criteria_text = generate_criteria_from_pdf(tmp_path)
        os.remove(tmp_path)

        return JSONResponse(content={"criteria": criteria_text})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/generate-criteria-sheet")
async def generate_criteria_sheet_endpoint(file: UploadFile = File(...)):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(await file.read())
            pdf_path = tmp.name

        # Étapes
        text = extract_text_from_pdf(pdf_path)
        criteria = generate_criteria_from_job_description(text)

        # Création fichier Word
        output_docx_path = os.path.join(tempfile.gettempdir(), "fiche_evaluation.docx")
        generate_docx(
            candidate_name="Goldman Sachs",
            job_title="Développeur Backend",
            analyst_name="Grace Esther",
            criteria_list=criteria,
            output_path=output_docx_path
        )

        os.remove(pdf_path)
        return FileResponse(output_docx_path, filename="fiche_evaluation.docx",
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002)